# src/core/file_handlers/doc_handler.py

"""
Enhanced DOC/DOCX handler with hybrid OCR support.
Integrates with your ocr_config.yaml for optimal performance.
"""

import os
import sys
import logging
import tempfile
import subprocess
import shutil
from pathlib import Path
from typing import List, Optional
import torch
from transformers import TrOCRProcessor, VisionEncoderDecoderModel, AutoTokenizer, AutoModel
import asyncio

import cv2
import numpy as np
from PIL import Image
import io
from docx import Document
import docx2txt

# Import the PDFHandler for processing converted PDFs.
from .pdf_handler import PDFHandler

# Ensure the project root is in the path if needed.
sys.path.insert(0, str(Path(__file__).parent.parent))

from src.core.ocr.tesseract_wrapper import TesseractOCR
from src.core.file_handlers.base_handler import FileHandler
from src.core.config import load_ocr_config

# Import the hybrid OCR system
try:
    from ..services.ocr_service import get_ocr_service
    from ..utils.config_loader import ConfigLoader
    HYBRID_OCR_AVAILABLE = True
except ImportError:
    HYBRID_OCR_AVAILABLE = False

logger = logging.getLogger(__name__)
logging.basicConfig(level=logging.DEBUG)

# Set PDFMiner logger to WARNING to hide debug messages
logging.getLogger("pdfminer").setLevel(logging.WARNING)

class AdvancedDocHandler(FileHandler):
    """
    Enhanced DOC/DOCX handler with hybrid OCR and smart processing.
    """

    def __init__(self, model_manager=None, use_smart_ocr=True):
        # Your existing setup
        self.ocr = TesseractOCR()
        self.pdf_handler = PDFHandler(model_manager=model_manager) if model_manager else PDFHandler()
        self.temp_dir = tempfile.TemporaryDirectory()

        self.model_manager = model_manager
        if model_manager:
            self.device = model_manager.get_device()
            self.handwritten_processor = model_manager.get_trocr_processor()
            self.handwritten_model = model_manager.get_trocr_model()
            self.bert_tokenizer = model_manager.get_klue_tokenizer()
            self.bert_model = model_manager.get_klue_bert()
        else:
            self.device = None
            self.handwritten_processor = None
            self.handwritten_model = None
            self.bert_tokenizer = None
            self.bert_model = None

        # Hybrid OCR integration
        self.use_smart_ocr = use_smart_ocr and HYBRID_OCR_AVAILABLE
        if self.use_smart_ocr:
            try:
                self.ocr_service = get_ocr_service()
                self.ocr_config = ConfigLoader.load_ocr_config()
                config_status = self.ocr_service.get_config_status()
                self.smart_ocr_available = config_status["tesseract_available"] or config_status["easyocr_available"]
                logger.info(f"DOC Handler smart OCR enabled: {self.smart_ocr_available}")
            except Exception as e:
                logger.warning(f"Smart OCR initialization failed: {e}, falling back to legacy mode")
                self.use_smart_ocr = False
                self.smart_ocr_available = False

    async def extract_text(self, file_path: str) -> str:
        """Extract text from a DOC or DOCX file with smart OCR strategy."""
        ext = Path(file_path).suffix.lower()
        if ext == '.docx':
            return await self._process_docx(file_path)
        elif ext == '.doc':
            return await self._process_doc(file_path)
        else:
            logger.error(f"Unsupported file format: {file_path}")
            return ""

    async def _process_docx(self, file_path: str) -> str:
        """Process a DOCX file with smart OCR strategy."""
        try:
            # Phase 1: Extract regular text content
            text_content = await asyncio.to_thread(self._parse_docx_document, file_path)
            total_text_length = len(text_content)
            
            # Phase 2: Determine OCR strategy for images
            should_process_images, ocr_mode = self._determine_smart_ocr_strategy(
                total_text_length, 1, "docx"
            )
            
            if should_process_images and self.use_smart_ocr and self.smart_ocr_available:
                logger.info(f"DOCX has {total_text_length} chars, using smart OCR mode: {ocr_mode}")
                image_texts = await self._extract_docx_images_smart(file_path, ocr_mode)
            else:
                # Use your original method
                image_texts = await self._extract_docx_images(file_path)
            
            # Combine text
            if image_texts:
                combined_text = f"{text_content}\n\n=== IMAGE TEXTS ===\n" + "\n".join(image_texts)
            else:
                combined_text = text_content
                
            return combined_text
            
        except Exception as e:
            logger.error(f"DOCX processing failed: {e}")
            return ""

    async def _extract_docx_images_smart(self, file_path: str, ocr_mode: str) -> List[str]:
        """Extract images from DOCX with smart OCR strategy."""
        if not self.use_smart_ocr:
            return await self._extract_docx_images(file_path)  # Fallback to original
            
        image_texts = []
        try:
            # Extract images to directory
            image_dir = await asyncio.to_thread(self._extract_images_to_dir, file_path)
            
            # Convert image files to PIL Images
            images = []
            image_paths = []
            for img_path in image_dir.glob("*"):
                if img_path.suffix.lower() in ['.png', '.jpg', '.jpeg']:
                    try:
                        img = Image.open(img_path)
                        images.append(img)
                        image_paths.append(img_path)
                    except Exception as e:
                        logger.warning(f"Failed to open image {img_path}: {e}")
                        continue
            
            if images:
                # Determine context for better OCR
                context = self._determine_image_context_docx(file_path, images)
                
                # Choose OCR engine
                if context == "handwritten" and self.handwritten_model:
                    # Use your TrOCR for handwritten content
                    image_texts = await self._extract_with_trocr_docx(images, image_paths)
                else:
                    # Use smart OCR service
                    engine = self._choose_ocr_engine_for_context(context)
                    extracted_texts = await self.ocr_service.extract_text_from_images(
                        images, ocr_mode, context, engine
                    )
                    image_texts = extracted_texts
                    
                logger.debug(f"Smart OCR processed {len(image_texts)} images from DOCX")
                
        except Exception as e:
            logger.error(f"Smart image extraction from DOCX failed: {e}")
            # Fallback to original method
            return await self._extract_docx_images(file_path)
            
        return image_texts

    async def _extract_with_trocr_docx(self, images: List[Image.Image], image_paths: List[Path]) -> List[str]:
        """Extract text using your existing TrOCR model for handwritten content."""
        image_texts = []
        
        for idx, (image, img_path) in enumerate(zip(images, image_paths)):
            try:
                # Convert PIL Image to the format your TrOCR expects
                image_bytes = io.BytesIO()
                image.save(image_bytes, format='PNG')
                image_bytes = image_bytes.getvalue()
                
                # Use your existing OCR processing
                text = await asyncio.to_thread(self._ocr_image_sync, str(img_path))
                if text:
                    image_texts.append(f"[TrOCR Image {idx + 1}] {text}")
                    
            except Exception as e:
                logger.warning(f"TrOCR failed for image {idx + 1}: {e}")
                continue
                
        return image_texts

    def _determine_image_context_docx(self, file_path: str, images: List[Image.Image]) -> str:
        """Determine the context of images in DOCX for optimal OCR."""
        try:
            # Analyze document content for context clues
            doc = Document(file_path)
            doc_text = ""
            for paragraph in doc.paragraphs:
                doc_text += paragraph.text.lower() + " "
                
            # Check for specific contexts
            if any(keyword in doc_text for keyword in ['handwritten', '손글씨', '필기', 'note', 'signature']):
                return "handwritten"
            elif any(keyword in doc_text for keyword in ['error', 'exception', 'traceback', '오류']):
                return "error_message"
            elif any(keyword in doc_text for keyword in ['diagram', 'figure', 'chart', '도표', '그림']):
                return "diagram"
            elif any(keyword in doc_text for keyword in ['screenshot', '스크린샷', 'capture']):
                return "screenshot"
            elif any(keyword in doc_text for keyword in ['한글', '한국어', '가나다']):
                return "korean_text"
            else:
                return "general"
                
        except Exception:
            return "general"

    def _determine_smart_ocr_strategy(self, text_length: int, page_count: int, file_type: str) -> tuple[bool, str]:
        """Determine if smart OCR should be used."""
        if not self.use_smart_ocr or not self.smart_ocr_available:
            # Use original logic
            return text_length < 1000, "full_legacy"
            
        # Smart logic for DOCX files
        if file_type == "docx":
            if text_length < 500:
                return True, "full"
            elif text_length < 2000:
                return True, "selective"
            elif text_length < 5000:
                return True, "minimal"
            else:
                return False, "skip"
        else:
            # Default logic
            return text_length < 1000, "selective"

    def _choose_ocr_engine_for_context(self, context: str) -> str:
        """Choose OCR engine based on context."""
        engine_preferences = {
            "handwritten": "auto",         # Let TrOCR handle this
            "korean_text": "easyocr",      # EasyOCR better for Korean
            "error_message": "tesseract",   # Tesseract good for clean text
            "screenshot": "tesseract",      # Screenshots usually have clean text
            "diagram": "auto",
            "general": "auto"
        }
        
        return engine_preferences.get(context, "auto")

    # Keep all your existing methods for backward compatibility
    async def extract_tables(self, file_path: str) -> List[List[List[str]]]:
        """Extract tables from a DOC or DOCX file."""
        ext = Path(file_path).suffix.lower()
        if ext == '.docx':
            return await self._extract_docx_tables(file_path)
        elif ext == '.doc':
            pdf_path = await self._convert_doc_to_pdf(file_path)
            if pdf_path and pdf_path.exists():
                try:
                    tables = await self.pdf_handler.extract_tables(str(pdf_path))
                    return tables
                finally:
                    pdf_path.unlink()
            else:
                logger.error("PDF conversion failed for table extraction.")
                return []
        else:
            logger.error(f"Unsupported file format: {file_path}")
            return []

    async def _process_doc(self, file_path: str) -> str:
        """Process a legacy .doc file."""
        try:
            pdf_path = await self._convert_doc_to_pdf(file_path)
            if pdf_path and pdf_path.exists():
                logger.debug(f"DOC converted to PDF: {pdf_path}")
                text = await self.pdf_handler.extract_text(str(pdf_path))
                pdf_path.unlink()  # Clean up the temporary PDF.
                return text
            else:
                raise Exception("PDF conversion returned no file.")
        except Exception as e:
            logger.error(f"LibreOffice PDF conversion failed: {e}")
            # Fallback to antiword extraction.
            text = await self._extract_via_antiword(file_path)
            return text

    async def _convert_doc_to_pdf(self, file_path: str) -> Optional[Path]:
        """Convert a DOC file to PDF using LibreOffice's command-line interface."""
        try:
            # Load configuration.
            config = load_ocr_config()
            libreoffice_config = config.get('libreoffice', {})
            soffice_path = libreoffice_config.get('path', None)

            if not soffice_path or not os.path.exists(soffice_path):
                logger.error("LibreOffice path not found in config or does not exist. Please check your config file.")
                return None

            output_dir = Path(self.temp_dir.name)
            cmd = [
                soffice_path,
                "--headless",
                "--convert-to",
                "pdf",
                "--outdir",
                str(output_dir),
                str(file_path)
            ]
            logger.debug(f"Running LibreOffice conversion command: {' '.join(cmd)}")
            
            # Run the command asynchronously
            proc = await asyncio.create_subprocess_exec(
                *cmd,
                stdout=asyncio.subprocess.PIPE,
                stderr=asyncio.subprocess.PIPE
            )
            await proc.communicate()
            
            pdf_file = output_dir / (Path(file_path).stem + ".pdf")
            if pdf_file.exists():
                logger.debug(f"PDF conversion successful: {pdf_file}")
                return pdf_file
            else:
                raise Exception("PDF file was not created by LibreOffice conversion.")
        except Exception as e:
            logger.error(f"LibreOffice conversion error: {e}")
            return None

    async def _extract_via_antiword(self, file_path: str) -> str:
        """Fallback text extraction for DOC files using antiword."""
        try:
            if not shutil.which("antiword"):
                logger.warning("antiword not found in system PATH. Please install antiword for fallback DOC extraction.")
                return ""
                
            # Use asyncio.subprocess for non-blocking operation
            proc = await asyncio.create_subprocess_exec(
                'antiword', '-m', 'UTF-8.txt', file_path,
                stdout=asyncio.subprocess.PIPE,
                stderr=asyncio.subprocess.PIPE,
                text=True
            )
            stdout, stderr = await proc.communicate()
            
            if proc.returncode == 0:
                logger.debug("Antiword extraction succeeded.")
                return stdout
            else:
                logger.warning(f"Antiword extraction failed with code {proc.returncode}: {stderr}")
                return ""
        except Exception as e:
            logger.warning(f"Antiword extraction failed: {e}")
            return ""

    def _parse_docx_document(self, file_path: str) -> str:
        """Opens a docx file and parses its structure."""
        doc = Document(file_path)
        return self._parse_docx_structure(doc)

    def _parse_docx_structure(self, doc: Document) -> str:
        """Parse the structure of a DOCX document."""
        content = []
        # Process paragraphs.
        for para in doc.paragraphs:
            indent = para.paragraph_format.left_indent or 0
            indent_str = " " * int(indent / 360)  # Approximate conversion from EMU to spaces.
            if para.style.name.startswith('List'):
                content.append(f"{indent_str}• {para.text}")
            else:
                content.append(f"{indent_str}{para.text}")
        # Process tables.
        for table in doc.tables:
            table_content = ["\n=== TABLE ==="]
            for row in table.rows:
                cells = [cell.text.replace('\n', ' ') for cell in row.cells]
                table_content.append(" | ".join(cells))
            content.append("\n".join(table_content))
        return "\n".join(content)

    async def _extract_docx_tables(self, file_path: str) -> List[List[List[str]]]:
        """Extract tables from a DOCX file."""
        try:
            return await asyncio.to_thread(self._extract_docx_tables_sync, file_path)
        except Exception as e:
            logger.error(f"Error extracting DOCX tables: {e}")
            return []
            
    def _extract_docx_tables_sync(self, file_path: str) -> List[List[List[str]]]:
        """Synchronous version of table extraction to be run in a thread pool."""
        tables_data = []
        doc = Document(file_path)
        for table in doc.tables:
            table_rows = []
            for row in table.rows:
                row_data = [cell.text.replace('\n', ' ') for cell in row.cells]
                table_rows.append(row_data)
            tables_data.append(table_rows)
        return tables_data

    async def _extract_docx_images(self, file_path: str) -> List[str]:
        """Extract images from a DOCX file and perform OCR (original method)."""
        image_texts = []
        try:
            # Run image extraction in a thread
            image_dir = await asyncio.to_thread(self._extract_images_to_dir, file_path)
            
            # Process each image file
            ocr_tasks = []
            for img_path in image_dir.glob("*"):
                if img_path.suffix.lower() in ['.png', '.jpg', '.jpeg']:
                    ocr_tasks.append(self._ocr_image(str(img_path)))
                    
            # Process images concurrently
            if ocr_tasks:
                image_texts = await asyncio.gather(*ocr_tasks)
                # Filter out empty results
                image_texts = [text for text in image_texts if text]
                
        except Exception as e:
            logger.error(f"Image extraction from DOCX failed: {e}")
        return image_texts
        
    def _extract_images_to_dir(self, file_path: str) -> Path:
        """Extract images from docx file to a directory. Returns the directory path."""
        image_dir = Path(self.temp_dir.name) / "images"
        image_dir.mkdir(exist_ok=True)
        # Extract images; docx2txt saves images in the specified folder
        docx2txt.process(file_path, str(image_dir))
        return image_dir

    async def _ocr_image(self, image_path: str) -> str:
        """Process an image with OCR to extract text."""
        try:
            # Run CPU-intensive image processing in a thread
            return await asyncio.to_thread(self._ocr_image_sync, image_path)
        except Exception as e:
            logger.warning(f"OCR failed for {image_path}: {e}")
            return ""
            
    def _ocr_image_sync(self, image_path: str) -> str:
        """Synchronous version of OCR processing to be run in a thread pool."""
        img = cv2.imread(image_path)
        if img is None:
            logger.warning(f"Unable to read image: {image_path}")
            return ""
            
        gray = cv2.cvtColor(img, cv2.COLOR_BGR2GRAY)
        thresh = cv2.adaptiveThreshold(gray, 255, cv2.ADAPTIVE_THRESH_GAUSSIAN_C, cv2.THRESH_BINARY, 11, 2)
        
        # Use TrOCR if available, otherwise fallback to basic OCR
        if self.handwritten_model and self.handwritten_processor:
            from PIL import Image
            pil_img = Image.fromarray(thresh)
            with torch.no_grad():
                inputs = self.handwritten_processor(images=pil_img, return_tensors="pt").to(self.handwritten_model.device)
                generated_ids = self.handwritten_model.generate(inputs.pixel_values)
                ocr_text = self.handwritten_processor.batch_decode(generated_ids, skip_special_tokens=True)[0]
            return ocr_text
        else:
            # Fallback to Tesseract if model is not available
            return self.ocr.ocr_image(thresh)

    def get_handler_status(self) -> dict:
        """Get current handler status for debugging."""
        status = {
            "handler": "Enhanced DOC Handler",
            "smart_ocr_enabled": self.use_smart_ocr,
            "smart_ocr_available": getattr(self, 'smart_ocr_available', False),
            "ml_models_loaded": {
                "trocr": self.handwritten_model is not None,
                "bert": self.bert_model is not None,
            }
        }
        
        if self.use_smart_ocr and hasattr(self, 'ocr_service'):
            status["ocr_service_status"] = self.ocr_service.get_config_status()
        
        return status